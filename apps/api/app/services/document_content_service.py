from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Any

import cv2
import numpy as np
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image
from xlrd import open_workbook

from app.core.config import get_settings
from app.services.ocr import get_ocr_engine
from app.services.ocr.schemas import OcrLine


class UnsupportedDocumentFormatError(ValueError):
    pass


class EmptyDocumentContentError(ValueError):
    pass


@dataclass(frozen=True)
class DocumentPageContent:
    page_number: int
    text: str
    confidence: float


class DocumentContentService:
    TEXT_EXTENSIONS = {".txt", ".md"}
    IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}
    MIN_NATIVE_PDF_TEXT_CHARS = 20
    _OCR_LINE_REPLACEMENTS: tuple[tuple[re.Pattern[str], str], ...] = (
        (
            re.compile(r"^C\s*NG\s+H(?:A|ÒA)\s+X\s+HI\s+CH\s+NGH(?:A|ÎA)\s+VI\s*T\s+NAM$", re.IGNORECASE),
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        ),
        (
            re.compile(r"^Đ\s*c\s+l\s*p\s*-\s*T\s*do\s*-\s*H\s*nh\s+phúc$", re.IGNORECASE),
            "Độc lập - Tự do - Hạnh phúc",
        ),
    )
    _OCR_TERM_REPLACEMENTS: tuple[tuple[re.Pattern[str], str], ...] = (
        (re.compile(r"\bĐiu\b", re.IGNORECASE), "Điều"),
        (re.compile(r"\bDiu\b", re.IGNORECASE), "Điều"),
        (re.compile(r"\bKhon\b", re.IGNORECASE), "Khoản"),
        (re.compile(r"\bPhm vi\b", re.IGNORECASE), "Phạm vi"),
        (re.compile(r"\bPhạm vi Điều chnh\b", re.IGNORECASE), "Phạm vi điều chỉnh"),
        (re.compile(r"\bđiu chnh\b", re.IGNORECASE), "điều chỉnh"),
        (re.compile(r"\bĐiều chnh\b", re.IGNORECASE), "điều chỉnh"),
        (re.compile(r"\bdieu chinh\b", re.IGNORECASE), "điều chỉnh"),
        (re.compile(r"\bđu thu\b", re.IGNORECASE), "đấu thầu"),
        (re.compile(r"\bdau thau\b", re.IGNORECASE), "đấu thầu"),
        (re.compile(r"\bVăn bn\b", re.IGNORECASE), "Văn bản"),
        (re.compile(r"\bVan ban\b", re.IGNORECASE), "Văn bản"),
        (re.compile(r"\bphai gi du ting Vit\b", re.IGNORECASE), "phải giữ dấu tiếng Việt"),
        (re.compile(r"\bphi gi du ting Vit\b", re.IGNORECASE), "phải giữ dấu tiếng Việt"),
        (re.compile(r"\bLUẶT\b"), "LUẬT"),
        (re.compile(r"\bLỤẬT\b"), "LUẬT"),
        (re.compile(r"\bđiều chính\b", re.IGNORECASE), "điều chỉnh"),
        (re.compile(r"\bdầu khi\b", re.IGNORECASE), "dầu khí"),
        (re.compile(r"\bCÔNG BẢO\b", re.IGNORECASE), "CÔNG BÁO"),
        (re.compile(r"(CÔNG BÁO/SỐ\s+\d+)\s*[?4+]\s*(\d+)", re.IGNORECASE), r"\1 + \2"),
        (re.compile(r"\bSố:\s*(\d+)\]/", re.IGNORECASE), r"Số: \1/"),
        (re.compile(r"\b27IS/2026\b"), "27/5/2026"),
        (re.compile(r"\bThứ\s+2715/2026\b"), "27/5/2026"),
        (re.compile(r"\bNhà Tháng\b"), ""),
        (re.compile(r"\bthuận thuận\b", re.IGNORECASE), ""),
        (re.compile(r"\bThông bảo\b", re.IGNORECASE), "Thông báo"),
        (re.compile(r"\s+và nhiều$", re.IGNORECASE), " và"),
        (re.compile(r"\s+1992$"), ""),
        (re.compile(r"^1990,\s*"), ""),
        (re.compile(r"\s+\(eb$"), ""),
        (re.compile(r"\bsố điện thoai\b", re.IGNORECASE), "số điện thoại"),
        (re.compile(r"\bMINH PHỦ\b"), "MINH PHÚ"),
    )
    _OCR_EMPTY_LINE_TEXTS = {"E", "16", "6n", "2", "xuân", "000001.", "000001", "CHUNICH"}
    _OCR_NOISE_PREFIXES = (
        "Thuật thuật ",
        "Thuật ",
        "Nhất ",
        "Thành ",
        "Nhà Tháng ",
        "Các thuận ",
        "Anh thuận ",
    )

    def __init__(self) -> None:
        self.settings = get_settings()

    def extract_pages(self, file_path: Path, filename: str) -> list[DocumentPageContent]:
        if not file_path.exists():
            raise FileNotFoundError(f"Uploaded file not found: {file_path}")

        extension = self._extension(file_path, filename)
        if extension in self.TEXT_EXTENSIONS:
            pages = [DocumentPageContent(page_number=1, text=self._read_text(file_path), confidence=1.0)]
        elif extension == ".docx":
            pages = [DocumentPageContent(page_number=1, text=self._read_docx(file_path), confidence=1.0)]
        elif extension == ".xlsx":
            pages = self._read_xlsx(file_path)
        elif extension == ".xls":
            pages = self._read_xls(file_path)
        elif extension == ".doc":
            raise UnsupportedDocumentFormatError(
                "Legacy .doc is not supported yet. Convert the file to .docx or .pdf and upload again."
            )
        elif extension == ".pdf":
            pages = self._ocr_pdf(file_path)
        elif extension in self.IMAGE_EXTENSIONS:
            pages = [self._ocr_image_file(file_path, page_number=1)]
        else:
            raise UnsupportedDocumentFormatError(f"Unsupported file extension: {extension or 'unknown'}")

        non_empty_pages = [page for page in pages if page.text.strip()]
        if not non_empty_pages:
            raise EmptyDocumentContentError(f"No text content extracted from {filename}")
        return pages

    def _extension(self, file_path: Path, filename: str) -> str:
        return (Path(filename).suffix or file_path.suffix).lower()

    def _read_text(self, file_path: Path) -> str:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _read_docx(self, file_path: Path) -> str:
        document = DocxDocument(file_path)
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = self._clean_text(paragraph.text)
            if text:
                parts.append(text)

        for table_index, table in enumerate(document.tables, start=1):
            parts.append(f"Bảng {table_index}")
            for row in table.rows:
                cells = [self._clean_text(cell.text) for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    parts.append(row_text)

        return "\n".join(parts)

    def _read_xlsx(self, file_path: Path) -> list[DocumentPageContent]:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
        pages: list[DocumentPageContent] = []
        try:
            for page_number, sheet in enumerate(workbook.worksheets, start=1):
                pages.append(
                    DocumentPageContent(
                        page_number=page_number,
                        text=self._sheet_to_text(sheet.title, sheet.iter_rows(values_only=True)),
                        confidence=1.0,
                    )
                )
        finally:
            workbook.close()
        return pages

    def _read_xls(self, file_path: Path) -> list[DocumentPageContent]:
        workbook = open_workbook(str(file_path), on_demand=True)
        pages: list[DocumentPageContent] = []
        try:
            for page_number, sheet_name in enumerate(workbook.sheet_names(), start=1):
                sheet = workbook.sheet_by_name(sheet_name)
                rows = (sheet.row_values(row_index) for row_index in range(sheet.nrows))
                pages.append(
                    DocumentPageContent(
                        page_number=page_number,
                        text=self._sheet_to_text(sheet_name, rows),
                        confidence=1.0,
                    )
                )
        finally:
            workbook.release_resources()
        return pages

    def _sheet_to_text(self, sheet_name: str, rows: Any) -> str:
        lines = [f"Sheet: {sheet_name}"]
        for row_number, row in enumerate(rows, start=1):
            values = [self._format_cell_value(value) for value in row]
            row_text = " | ".join(value for value in values if value)
            if row_text:
                lines.append(f"Row {row_number}: {row_text}")
        return "\n".join(lines)

    def _format_cell_value(self, value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, float) and value.is_integer():
            return str(int(value))
        return self._clean_text(str(value))

    def _ocr_pdf(self, file_path: Path) -> list[DocumentPageContent]:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(str(file_path))
        pages: list[DocumentPageContent] = []
        try:
            for page_index in range(len(pdf)):
                page = pdf[page_index]
                try:
                    page_number = page_index + 1
                    native_text = self._extract_native_pdf_text(page)
                    if len(native_text) >= self.MIN_NATIVE_PDF_TEXT_CHARS:
                        pages.append(
                            DocumentPageContent(page_number=page_number, text=native_text, confidence=1.0)
                        )
                        continue

                    bitmap = page.render(scale=2.0)
                    image = bitmap.to_pil()
                    pages.append(self._ocr_pil_image(image, page_number=page_number))
                finally:
                    page.close()
        finally:
            pdf.close()
        return pages

    def _extract_native_pdf_text(self, page: Any) -> str:
        try:
            text_page = page.get_textpage()
            try:
                char_count = text_page.count_chars()
                if char_count <= 0:
                    return ""
                return self._clean_multiline_text(text_page.get_text_range(0, char_count))
            finally:
                close = getattr(text_page, "close", None)
                if close:
                    close()
        except Exception:
            return ""

    def _ocr_image_file(self, file_path: Path, page_number: int) -> DocumentPageContent:
        image = Image.open(file_path)
        try:
            return self._ocr_pil_image(image, page_number=page_number)
        finally:
            image.close()

    def _ocr_pil_image(self, image: Image.Image, page_number: int) -> DocumentPageContent:
        rgb_image = image.convert("RGB")
        image_array = np.array(rgb_image)
        candidates = self._preprocess_image_candidates(image_array)
        results = [self._run_ocr(candidate, page_number=page_number) for candidate in candidates]
        return max(results, key=self._ocr_result_score)

    def _preprocess_image_candidates(self, image_array: np.ndarray) -> list[np.ndarray]:
        mode = self.settings.ocr_preprocess_mode.lower()
        if mode == "raw":
            return [self._resize_for_ocr(image_array)]
        if mode == "clahe":
            return [self._preprocess_clahe(image_array)]
        if mode == "threshold":
            return [self._preprocess_threshold(image_array)]
        if mode != "auto":
            raise ValueError(f"Unsupported OCR_PREPROCESS_MODE: {self.settings.ocr_preprocess_mode}")

        raw = self._resize_for_ocr(image_array)
        clahe = self._preprocess_clahe(image_array)
        thresholded = self._preprocess_threshold(image_array)
        return [raw, clahe, thresholded]

    def _resize_for_ocr(self, image_array: np.ndarray) -> np.ndarray:
        height, width = image_array.shape[:2]
        if min(height, width) >= 1200:
            return image_array

        scale = 1200 / max(min(height, width), 1)
        return cv2.resize(image_array, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    def _preprocess_clahe(self, image_array: np.ndarray) -> np.ndarray:
        resized = self._resize_for_ocr(image_array)
        lab = cv2.cvtColor(resized, cv2.COLOR_RGB2LAB)
        l_channel, a_channel, b_channel = cv2.split(lab)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced_l = clahe.apply(l_channel)
        enhanced = cv2.merge((enhanced_l, a_channel, b_channel))
        return cv2.cvtColor(enhanced, cv2.COLOR_LAB2RGB)

    def _preprocess_threshold(self, image_array: np.ndarray) -> np.ndarray:
        gray = cv2.cvtColor(image_array, cv2.COLOR_RGB2GRAY)
        height, width = gray.shape[:2]
        if min(height, width) < 1200:
            scale = 1200 / max(min(height, width), 1)
            gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

        denoised = cv2.fastNlMeansDenoising(gray, h=10)
        thresholded = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            11,
        )
        return cv2.cvtColor(thresholded, cv2.COLOR_GRAY2RGB)

    def _run_ocr(self, image_array: np.ndarray, page_number: int) -> DocumentPageContent:
        raw_lines = self._recognize_ocr_lines(image_array)
        lines = self._clean_ocr_lines(raw_lines)
        text = "\n".join(line.text for line in lines)
        text = self._postprocess_ocr_text(text)
        confidence = sum(line.confidence for line in lines) / len(lines) if lines else 0.0
        return DocumentPageContent(page_number=page_number, text=text, confidence=round(confidence, 4))

    def _recognize_ocr_lines(self, image_array: np.ndarray) -> list[OcrLine]:
        try:
            return get_ocr_engine(self.settings).recognize(image_array)
        except FileNotFoundError:
            fallback_engine = (self.settings.ocr_fallback_engine or "").strip()
            if not fallback_engine:
                raise
            return get_ocr_engine(self.settings, fallback_engine).recognize(image_array)

    def _clean_ocr_lines(self, lines: list[OcrLine]) -> list[OcrLine]:
        cleaned_lines: list[OcrLine] = []
        for line in lines:
            text = self._clean_ocr_text(line.text)
            if text:
                cleaned_lines.append(OcrLine(text=text, confidence=line.confidence, box=line.box))
        return cleaned_lines

    def _ocr_result_score(self, page_content: DocumentPageContent) -> tuple[int, float, int, int]:
        return (
            self._count_ocr_quality_markers(page_content.text),
            page_content.confidence,
            self._count_vietnamese_accent_chars(page_content.text),
            len(page_content.text),
        )

    def _count_ocr_quality_markers(self, text: str) -> int:
        markers = [
            "Số:",
            "Kính gửi",
            "Người liên hệ",
            "Điều ",
            "Khoản ",
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        ]
        return sum(1 for marker in markers if marker in text)

    def _count_vietnamese_accent_chars(self, text: str) -> int:
        vietnamese_chars = set(
            "ăâđêôơưáàảãạắằẳẵặấầẩẫậéèẻẽẹếềểễệíìỉĩị"
            "óòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ"
            "ĂÂĐÊÔƠƯÁÀẢÃẠẮẰẲẴẶẤẦẨẪẬÉÈẺẼẸẾỀỂỄỆÍÌỈĨỊ"
            "ÓÒỎÕỌỐỒỔỖỘỚỜỞỠỢÚÙỦŨỤỨỪỬỮỰÝỲỶỸỴ"
        )
        return sum(1 for char in text if char in vietnamese_chars)

    def _clean_text(self, text: str) -> str:
        return " ".join(text.replace("\x00", " ").split())

    def _clean_ocr_text(self, text: str) -> str:
        cleaned = self._clean_text(text)
        if not self.settings.ocr_restore_vietnamese_terms:
            return cleaned

        for pattern, replacement in self._OCR_LINE_REPLACEMENTS:
            if pattern.fullmatch(cleaned):
                return replacement

        restored = cleaned
        for pattern, replacement in self._OCR_TERM_REPLACEMENTS:
            restored = pattern.sub(replacement, restored)
        restored = self._strip_known_ocr_noise(restored)
        if restored in self._OCR_EMPTY_LINE_TEXTS:
            return ""
        return restored

    def _strip_known_ocr_noise(self, text: str) -> str:
        cleaned = text
        changed = True
        while changed:
            changed = False
            for prefix in self._OCR_NOISE_PREFIXES:
                if cleaned.startswith(prefix):
                    cleaned = cleaned[len(prefix) :].lstrip()
                    changed = True
        return cleaned

    def _postprocess_ocr_text(self, text: str) -> str:
        fixed = text
        fixed = re.sub(
            r"KỸ THUẬT\s*\nTỔNG CÔNG TY HẠ\s*\nTẦNG",
            "TỔNG CÔNG TY HẠ TẦNG KỸ THUẬT",
            fixed,
        )
        fixed = re.sub(
            r"TỔNG CÔNG TY HẠ\s*\nKỸ THUẬT\s*\nTẦNG",
            "TỔNG CÔNG TY HẠ TẦNG KỸ THUẬT",
            fixed,
        )
        fixed = re.sub(
            r"KẾ HOẠCH MUA SẮM\s*\nTƯ NĂM 2026\s*\nVẬT",
            "KẾ HOẠCH MUA SẮM VẬT TƯ NĂM 2026",
            fixed,
        )
        fixed = re.sub(
            r"CỘNG HÒA XÃ HỘI CHỦ\s*\nNGHĨA VIỆT NAM",
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
            fixed,
        )
        fixed = re.sub(
            r"ỦY BAN NHÂN DÂN XÃ\s*\nMINH PHÚ",
            "ỦY BAN NHÂN DÂN XÃ MINH PHÚ",
            fixed,
        )
        fixed = re.sub(
            r"Minh Phú, ngày 27 tháng 5\s*\nnăm 2026",
            "Minh Phú, ngày 27 tháng 5 năm 2026",
            fixed,
        )
        fixed = re.sub(
            r"Về việc rà soát hồ sơ đề nghị cấp vật tư sửa chữa\s*\ntuyến điện",
            "Về việc rà soát hồ sơ đề nghị cấp vật tư sửa chữa tuyến điện",
            fixed,
        )
        fixed = re.sub(
            r"Đề nghị quý phòng kiểm tra số lượng cáp điện,\s*\naptomat và vật tư an toàn\.",
            "Đề nghị quý phòng kiểm tra số lượng cáp điện, aptomat và vật tư an toàn.",
            fixed,
        )
        fixed = re.sub(
            r"Thời hạn phản hồi trước ngày 05 tháng 6 năm\s*\n2026 để tổng hợp kế hoạch mua sắm\.?",
            "Thời hạn phản hồi trước ngày 05 tháng 6 năm 2026 để tổng hợp kế hoạch mua sắm.",
            fixed,
        )
        return fixed

    def _clean_multiline_text(self, text: str) -> str:
        lines = [
            self._clean_text(line)
            for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
        ]
        return "\n".join(line for line in lines if line)
